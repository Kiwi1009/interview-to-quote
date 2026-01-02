import os
from typing import Optional
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session
from app.models.case import Case
from app.models.extraction_run import ExtractionRun
from app.models.plan import Plan
from app.models.quote_item import QuoteItem
from app.models.extracted_requirement import ExtractedRequirement
from app.models.document import DocumentType, DocumentFormat
from app.core.config import settings
from app.services.extraction_service import ExtractionService
from app.services.plan_service import PlanService

class DocumentGenerator:
    def __init__(self, db: Session):
        self.db = db
    
    def generate_document(self, case_id: int, run_id: Optional[int], doc_type: DocumentType, format: DocumentFormat) -> Optional[str]:
        case = self.db.query(Case).filter(Case.id == case_id).first()
        if not case:
            return None
        
        if doc_type == DocumentType.SPEC:
            return self._generate_spec_document(case, run_id, format)
        elif doc_type == DocumentType.REPORT:
            return self._generate_report_document(case, run_id, format)
        elif doc_type == DocumentType.QUOTE:
            return self._generate_quote_document(case, run_id, format)
        return None
    
    def _generate_spec_document(self, case: Case, run_id: Optional[int], format: DocumentFormat) -> Optional[str]:
        """Generate Requirements Specification document"""
        extraction_service = ExtractionService(self.db)
        requirements = extraction_service.get_requirements(case.id, run_id)
        if not requirements:
            return None
        
        doc = DocxDocument()
        
        # Title
        title = doc.add_heading('需求規格書', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Case info
        doc.add_heading('案件資訊', 1)
        doc.add_paragraph(f'案件名稱：{case.title}')
        if case.industry:
            doc.add_paragraph(f'產業別：{case.industry}')
        
        # Requirements
        req_data = requirements.jsonb_data
        doc.add_heading('需求內容', 1)
        
        if req_data.get('customer_pain_points'):
            doc.add_heading('客戶痛點', 2)
            for point in req_data['customer_pain_points']:
                doc.add_paragraph(point, style='List Bullet')
        
        if req_data.get('workpiece'):
            doc.add_heading('工件資訊', 2)
            wp = req_data['workpiece']
            for key, value in wp.items():
                if value:
                    doc.add_paragraph(f'{key}：{value}')
        
        if req_data.get('process'):
            doc.add_heading('製程資訊', 2)
            proc = req_data['process']
            for key, value in proc.items():
                if value:
                    doc.add_paragraph(f'{key}：{value}')
        
        if req_data.get('constraints'):
            doc.add_heading('限制條件', 2)
            constraints = req_data['constraints']
            for key, value in constraints.items():
                if value:
                    doc.add_paragraph(f'{key}：{value}')
        
        if req_data.get('open_questions'):
            doc.add_heading('開放問題', 2)
            for q in req_data['open_questions']:
                doc.add_paragraph(q, style='List Bullet')
        
        # Save
        filename = f"spec_{case.id}_{run_id or 'latest'}.docx"
        filepath = os.path.join(settings.DOCUMENT_PATH, filename)
        doc.save(filepath)
        
        if format == DocumentFormat.PDF:
            # Convert to PDF (requires additional library or use WeasyPrint)
            pdf_path = filepath.replace('.docx', '.pdf')
            # For now, return DOCX path - PDF conversion can be added
            return filepath
        
        return filepath
    
    def _generate_report_document(self, case: Case, run_id: Optional[int], format: DocumentFormat) -> Optional[str]:
        """Generate Requirements Report (table format)"""
        extraction_service = ExtractionService(self.db)
        requirements = extraction_service.get_requirements(case.id, run_id)
        if not requirements:
            return None
        
        doc = DocxDocument()
        doc.add_heading('需求報告表', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create table
        req_data = requirements.jsonb_data
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = '欄位'
        header_cells[1].text = '內容'
        header_cells[2].text = '證據'
        
        # Add rows
        def add_table_row(field_path: str, value: Any):
            row = table.add_row()
            row.cells[0].text = field_path
            row.cells[1].text = str(value) if value else 'N/A'
            # Find evidence
            evidence_text = ''
            for ev in requirements.evidence:
                if ev.field_path == field_path:
                    evidence_text = ev.snippet[:100] + '...' if len(ev.snippet) > 100 else ev.snippet
                    break
            row.cells[2].text = evidence_text
        
        # Add data rows
        if req_data.get('workpiece'):
            for key, value in req_data['workpiece'].items():
                add_table_row(f'workpiece.{key}', value)
        
        if req_data.get('process'):
            for key, value in req_data['process'].items():
                add_table_row(f'process.{key}', value)
        
        filename = f"report_{case.id}_{run_id or 'latest'}.docx"
        filepath = os.path.join(settings.DOCUMENT_PATH, filename)
        doc.save(filepath)
        
        return filepath
    
    def _generate_quote_document(self, case: Case, run_id: Optional[int], format: DocumentFormat) -> Optional[str]:
        """Generate Quotation document with 3 plans"""
        plan_service = PlanService(self.db)
        plans = plan_service.list_plans(case.id, run_id)
        if not plans:
            return None
        
        doc = DocxDocument()
        doc.add_heading('報價單', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'案件名稱：{case.title}')
        doc.add_paragraph(f'日期：{case.created_at.strftime("%Y-%m-%d")}')
        
        # Generate quote for each plan
        for plan in plans:
            doc.add_heading(f'{plan.plan_code.value} - {plan.name}', 1)
            
            # Quote items table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            header = table.rows[0].cells
            header[0].text = '類別'
            header[1].text = '項目名稱'
            header[2].text = '規格'
            header[3].text = '數量'
            header[4].text = '單價（低-高）'
            header[5].text = '小計（低-高）'
            
            total_low = 0
            total_high = 0
            
            for item in plan.quote_items:
                row = table.add_row()
                row.cells[0].text = item.category
                row.cells[1].text = item.item_name
                row.cells[2].text = item.spec or ''
                row.cells[3].text = f"{item.qty} {item.unit}"
                row.cells[4].text = f"${item.unit_price_low:,.0f} - ${item.unit_price_high:,.0f}"
                row.cells[5].text = f"${item.subtotal_low or 0:,.0f} - ${item.subtotal_high or 0:,.0f}"
                total_low += item.subtotal_low or 0
                total_high += item.subtotal_high or 0
            
            # Totals
            doc.add_paragraph(f'小計：${total_low:,.0f} - ${total_high:,.0f}')
            contingency = total_low * 0.1  # 10% contingency
            doc.add_paragraph(f'預備費（10%）：${contingency:,.0f}')
            doc.add_paragraph(f'合計：${total_low + contingency:,.0f} - ${total_high + contingency:,.0f}')
            doc.add_paragraph('')  # Spacing
        
        filename = f"quote_{case.id}_{run_id or 'latest'}.docx"
        filepath = os.path.join(settings.DOCUMENT_PATH, filename)
        doc.save(filepath)
        
        return filepath

